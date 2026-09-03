# -*- coding: utf-8 -*-
"""Sinh bao cao Word theo chuan Phu luc 2 khoa HTTT UIT.

Chuan hinh thuc: Times New Roman 13pt, gian dong 1.5, le tren 3cm / duoi 3.5cm /
trai 3.5cm / phai 2cm, so trang giua ben duoi, tieu de chuong bold 14pt, tieu de
muc bold 13pt, label hinh nam DUOI hinh, label bang nam TREN bang.

Noi dung: phan cua Danh da hoan thanh. Cac muc cua thanh vien khac de placeholder.
"""
import pathlib
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = pathlib.Path(__file__).resolve().parent.parent
OUT = ROOT / "report" / "bao-cao-quy-trinh-nghiep-vu-TGDD-TopZone.docx"
IMG = ROOT / "model" / "hinh-xuat"
FIG11 = ROOT / "docs" / "kien-truc-quy-trinh" / "hinh-1-1-kien-truc-quy-trinh.png"

FONT, SZ = "Times New Roman", 13


def set_font(run, size=SZ, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SZ)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for s in doc.sections:
        s.top_margin, s.bottom_margin = Cm(3), Cm(3.5)
        s.left_margin, s.right_margin = Cm(3.5), Cm(2)


def page_number(doc):
    """So trang giua ben duoi."""
    for s in doc.sections:
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        set_font(r)
        for el, attr in (("w:fldChar", {"w:fldCharType": "begin"}),
                         ("w:instrText", None), ("w:fldChar", {"w:fldCharType": "end"})):
            e = OxmlElement(el)
            if attr:
                for k, v in attr.items():
                    e.set(qn(k), v)
            else:
                e.set(qn("xml:space"), "preserve")
                e.text = " PAGE "
            r._r.append(e)


def h1(doc, text):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(12), Pt(12)
    set_font(p.add_run(text.upper()), 14, bold=True)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    set_font(p.add_run(text), SZ, bold=True)
    return p


def para(doc, text, italic=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    set_font(p.add_run(text), italic=italic)
    return p


def todo(doc, who, what):
    p = doc.add_paragraph()
    set_font(p.add_run("[CHUA CO] %s — %s" % (who, what)), italic=True,
             color=RGBColor(0xB0, 0x30, 0x30))
    return p


def figure(doc, path, caption):
    """Anh + label NAM DUOI theo chuan."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(15.5))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(12)
    set_font(c.add_run(caption), 12, italic=True)


def table_cap(doc, caption):
    """Label bang NAM TREN theo chuan."""
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(4)
    set_font(c.add_run(caption), 12, italic=True)


def grid(doc, rows, widths=None):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            set_font(p.add_run(str(v)), 12, bold=(i == 0))
    return t


def build():
    doc = Document()
    base_style(doc)
    page_number(doc)

    # ---------- Trang bia ----------
    for t, sz, b in [("ĐẠI HỌC QUỐC GIA THÀNH PHỐ HỒ CHÍ MINH", 13, True),
                     ("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN", 13, True),
                     ("KHOA HỆ THỐNG THÔNG TIN", 13, True), ("", 13, False),
                     ("BÁO CÁO ĐỒ ÁN MÔN HỌC", 14, True),
                     ("HỆ THỐNG QUẢN TRỊ QUI TRÌNH NGHIỆP VỤ", 13, True), ("", 13, False),
                     ("PHÂN TÍCH QUY TRÌNH NGHIỆP VỤ", 16, True),
                     ("CHUỖI BÁN LẺ THEGIOIDIDONG.COM VÀ TOPZONE", 16, True)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(t), sz, bold=b)
    doc.add_paragraph()
    table_cap(doc, "Bảng 0.1. Danh sách thành viên nhóm")
    grid(doc, [["MSSV", "Họ và tên", "Vai trò"],
               ["24730090", "Nguyễn Ngọc Danh", "Nhóm trưởng"],
               ["24730132", "Nguyễn Thị Hồng Phúc", "Thành viên"],
               ["24730131", "Nguyễn Thanh Phúc", "Thành viên"],
               ["24730099", "Mai Hoàng Hưng", "Thành viên"]])
    doc.add_paragraph()
    for t in ["Giảng viên hướng dẫn: ThS. Hà Lê Hoài Trung", "TP. Hồ Chí Minh, tháng 9 năm 2026"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(t))

    # ---------- Bon danh muc ----------
    for name in ["MỤC LỤC", "DANH MỤC HÌNH VẼ", "DANH MỤC BẢNG BIỂU", "DANH MỤC TỪ VIẾT TẮT"]:
        h1(doc, name)
        todo(doc, "Hồng Phúc", "sinh tự động sau khi nội dung các chương đã chốt (hạn 04/09)")

    # ---------- Chuong 1 ----------
    h1(doc, "Chương 1. Tổng quan doanh nghiệp và kiến trúc quy trình")
    h2(doc, "1.1. Giới thiệu doanh nghiệp")
    para(doc, "Công ty Cổ phần Đầu tư Thế Giới Di Động (MWG) vận hành hai chuỗi bán lẻ "
              "thiết bị công nghệ được chọn làm đối tượng phân tích: thegioididong.com — "
              "chuỗi phổ thông đa thương hiệu, và TopZone — chuỗi ủy quyền cao cấp của "
              "Apple. Hai chuỗi dùng chung nền tảng vận hành (ERP, hệ thống kho, hệ thống "
              "bảo hành), nên phân tích quy trình lõi có thể gộp và chỉ tách khi có khác "
              "biệt thật sự về nghiệp vụ.")
    todo(doc, "Danh", "bổ sung số liệu quy mô từ báo cáo thường niên MWG kèm nguồn và ngày truy cập")

    h2(doc, "1.2. Phạm vi đề tài và giới hạn")
    para(doc, "Phạm vi sản phẩm thu hẹp ở bốn nhóm: điện thoại di động, laptop, máy tính "
              "bảng và phụ kiện đi kèm. Không bao gồm đồng hồ thông minh, hàng gia dụng "
              "của Điện Máy Xanh, hàng tiêu dùng của Bách Hóa Xanh và các dịch vụ SIM số.")
    para(doc, "Nhóm là người quan sát bên ngoài, không có quyền truy cập hệ thống nội bộ. "
              "Do đó toàn bộ số liệu vận hành trong Chương 4 đến từ bấm giờ trực tiếp tại "
              "cửa hàng (mẫu nhỏ, có ghi rõ cỡ mẫu n) hoặc từ bảng giả định có đánh dấu "
              "\u201c(ước lượng)\u201d. Các quy trình mô tả trong báo cáo là mô hình suy "
              "luận từ quan sát, không phải quy trình chuẩn do MWG ban hành.")

    h2(doc, "1.3. Kiến trúc quy trình nghiệp vụ")
    para(doc, "Hoạt động của chuỗi được phân rã thành 12 quy trình, chia ba lớp: lớp cốt "
              "lõi trực tiếp sinh doanh thu, lớp quản lý điều phối nguồn lực và ra quyết "
              "định, lớp hỗ trợ duy trì năng lực vận hành. Trong đó sáu quy trình có độ "
              "phức tạp nghiệp vụ cao được chọn để mô hình hóa bằng BPMN.")
    figure(doc, FIG11, "Hình 1.1. Kiến trúc quy trình nghiệp vụ chuỗi TGDĐ và TopZone")
    table_cap(doc, "Bảng 1.1. Danh mục 12 quy trình theo ba lớp")
    grid(doc, [["Lớp", "Mã", "Tên quy trình", "BPMN"],
               ["Cốt lõi", "C1", "Bán tại cửa hàng", "—"],
               ["", "C2", "Bán online, giao hàng và nhận tại cửa hàng", "—"],
               ["", "C3", "Bán trả góp", "Có"],
               ["", "C4", "Bảo hành, đổi trả", "Có"],
               ["Quản lý", "M1", "Hoạch định nhu cầu", "—"],
               ["", "M2", "Quản lý nhà cung cấp", "Có"],
               ["", "M3", "Kho và điều chuyển", "Có"],
               ["", "M4", "Mạng lưới cửa hàng", "—"],
               ["Hỗ trợ", "S1", "Tuyển dụng và đào tạo", "Có"],
               ["", "S2", "Vận hành ERP / POS", "—"],
               ["", "S3", "Mua sắm hạ tầng", "—"],
               ["", "S4", "Đối soát công nợ nhà cung cấp", "Có"]])

    # ---------- Chuong 2 ----------
    h1(doc, "Chương 2. Cơ sở lý thuyết và phương pháp")
    for m, who, what in [("2.1. Khái niệm quy trình nghiệp vụ và BPM lifecycle", "Hồng Phúc", "nội dung lý thuyết"),
                         ("2.2. Ký hiệu BPMN 2.0 sử dụng trong đề tài", "Hồng Phúc", "bảng ký hiệu"),
                         ("2.3. Phân loại lãng phí Lean theo nhóm Move / Hold / Overdo", "Hồng Phúc", "khung phân loại"),
                         ("2.4. Phương pháp thu thập dữ liệu", "Hưng", "mô tả cách quan sát, bấm giờ, phỏng vấn")]:
        h2(doc, m)
        todo(doc, who, what)
    h2(doc, "2.5. Quy trình review chéo mô hình trong nhóm")
    para(doc, "Mỗi mô hình BPMN do một thành viên vẽ và một thành viên khác duyệt thông "
              "qua pull request trên GitHub. Comment trong pull request được dùng làm biên "
              "bản review chéo. Kết quả tự kiểm và nhận xét của người duyệt được tổng hợp "
              "trong hồ sơ bằng chứng của nhóm.")
    table_cap(doc, "Bảng 2.1. Kết quả tự kiểm ba mô hình do nhóm trưởng vẽ")
    grid(doc, [["Tiêu chí", "M2", "C3", "C4"], ["Số gateway (yêu cầu > 7)", "9", "10", "11"],
               ["Số lane", "5", "5", "6"], ["Số task", "15", "15", "16"],
               ["Số sequence flow", "35", "35", "41"],
               ["Luồng trỏ sai / node mồ côi", "0", "0", "0"]])
    todo(doc, "Hồng Phúc", "bổ sung nhận xét người duyệt sau khi review PR (hạn 30/08)")
    return doc


def chuong_3_4(doc):
    # ---------- Chuong 3 ----------
    h1(doc, "Chương 3. Mô tả và mô hình hóa quy trình")
    h2(doc, "3.1. Nhóm quy trình cốt lõi")
    para(doc, "Ba quy trình C1, C3 và C4 được lập hồ sơ đầy đủ theo template 12 mục thống "
              "nhất của nhóm: mục đích, phạm vi, actor, đầu vào và đầu ra, các bước thực "
              "hiện, điểm ra quyết định, ngoại lệ, quy tắc nghiệp vụ, chỉ số đo lường, hệ "
              "thống và biểu mẫu, điểm nghẽn quan sát được, nguồn tham chiếu.")
    table_cap(doc, "Bảng 3.1. Tổng hợp ba quy trình cốt lõi do nhóm trưởng phụ trách")
    grid(doc, [["Mã", "Tên quy trình", "Số bước", "Điểm quyết định", "Ngoại lệ", "Điểm nghẽn"],
               ["C1", "Bán tại cửa hàng", "14", "9", "6", "4"],
               ["C3", "Bán trả góp", "15", "10", "7", "4"],
               ["C4", "Bảo hành, đổi trả", "16", "11", "7", "5"]])
    para(doc, "Hồ sơ chi tiết của từng quy trình được đưa vào Phụ lục. Ở đây chỉ trình bày "
              "phần mô hình hóa và các điểm nghẽn phục vụ phân tích ở Chương 4.")
    h2(doc, "3.2. Nhóm quy trình quản lý")
    todo(doc, "Hồng Phúc", "hồ sơ M1, M2, M3, M4 (hạn 22/08)")
    h2(doc, "3.3. Nhóm quy trình hỗ trợ")
    todo(doc, "Thanh Phúc", "hồ sơ S1, S2, S3, S4 (hạn 22/08)")

    h2(doc, "3.4. Mô hình hóa quy trình bằng BPMN")
    para(doc, "Ba mô hình dưới đây do nhóm trưởng vẽ. Mỗi mô hình vượt mốc bảy gateway "
              "theo yêu cầu về độ phức tạp. File nguồn định dạng .bpmn chuẩn BPMN 2.0, mở "
              "được bằng bpmn.io hoặc Camunda Modeler.")
    for code, fn, cap, extra in [
            ("M2", "M2-quan-ly-nha-cung-cap.png",
             "Hình 3.1. Mô hình BPMN quy trình M2 — Quản lý nhà cung cấp (9 gateway)",
             "Mô hình M2 được dựng khi hồ sơ M2 chưa hoàn tất, cần đối chiếu lại trước mốc "
             "khóa mô hình ngày 30/08."),
            ("C3", "C3-ban-tra-gop.png",
             "Hình 3.2. Mô hình BPMN quy trình C3 — Bán trả góp (10 gateway)",
             "Điểm đáng chú ý là gateway G4 ba nhánh ứng với ba kết quả thẩm định tín dụng, "
             "và vòng lặp bổ sung hồ sơ quay lại khâu thẩm định."),
            ("C4", "C4-bao-hanh-doi-tra.png",
             "Hình 3.3. Mô hình BPMN quy trình C4 — Bảo hành, đổi trả (11 gateway)",
             "Trung tâm bảo hành được vẽ thành lane riêng vì là tổ chức bên ngoài; quyết "
             "định cuối về lỗi phần cứng thuộc về actor này chứ không thuộc cửa hàng.")]:
        pth = IMG / fn
        if pth.exists():
            figure(doc, pth, cap)
        para(doc, extra)

    # ---------- Chuong 4 ----------
    h1(doc, "Chương 4. Phân tích và đề xuất cải tiến")
    for m, who, what in [("4.1. Bảng giả định cho số liệu chưa xác minh", "Hưng", "bảng giả định (hạn 28/08)"),
                         ("4.2. Phân tích định lượng: cycle time và CTE", "Hưng", "số liệu từ buổi khảo sát (hạn 01/09)"),
                         ("4.3. Phân tích giá trị VA / BVA / NVA", "Hồng Phúc", "bảng VA/BVA/NVA cho C3 và C4 (hạn 03/09)"),
                         ("4.4. Nhóm lãng phí Move / Hold / Overdo", "Hồng Phúc", "bảng phân nhóm (hạn 04/09)"),
                         ("4.5. Biểu đồ xương cá cho điểm nghẽn chính", "Hồng Phúc", "fishbone (hạn 04/09)")]:
        h2(doc, m)
        todo(doc, who, what)

    h2(doc, "4.6. Issue Register — tổng hợp phát hiện")
    para(doc, "Mười ba phát hiện dưới đây được gom từ mục điểm nghẽn của ba hồ sơ quy "
              "trình cốt lõi. Cột mức độ và cột bằng chứng chỉ điền được sau buổi khảo sát "
              "thực địa, nên bảng chưa xếp thứ tự ưu tiên — xếp ưu tiên khi chưa có số đo "
              "sẽ tạo ra thứ tự sai.")
    table_cap(doc, "Bảng 4.1. Issue Register — 13 phát hiện từ C1, C3 và C4")
    rows = [["Mã", "Quy trình", "Phát hiện", "Nhóm lãng phí"]]
    for m, q, t, g in [
        ("IR-01", "C1", "Kích hoạt máy và chuyển dữ liệu chiếm phần lớn thời gian giao dịch", "Overdo"),
        ("IR-02", "C1", "Kiểm tra tồn kho diễn ra muộn, công tư vấn trước đó bị bỏ phí", "Overdo"),
        ("IR-03", "C1", "Khách chờ tới lượt được tư vấn trong giờ cao điểm", "Hold"),
        ("IR-04", "C1", "Nhiều giao dịch nghẽn ở một quầy thu ngân", "Hold"),
        ("IR-05", "C3", "Chờ thẩm định tín dụng từ actor bên ngoài", "Hold"),
        ("IR-06", "C3", "Vòng lặp bổ sung hồ sơ nhân đôi thời gian chờ", "Overdo"),
        ("IR-07", "C3", "Nhập liệu hồ sơ thủ công, dễ sai và phải làm lại", "Overdo"),
        ("IR-08", "C3", "Hết hàng sau khi hồ sơ đã duyệt", "Move"),
        ("IR-09", "C4", "Khách chờ tới lượt tại quầy bảo hành", "Hold"),
        ("IR-10", "C4", "Chờ kết quả từ trung tâm bảo hành — khoảng chờ dài nhất", "Hold"),
        ("IR-11", "C4", "Tra cứu thủ công khi khách không có hóa đơn", "Overdo"),
        ("IR-12", "C4", "Máy nhận về không đạt phải gửi lại", "Move + Overdo"),
        ("IR-13", "C4", "Vòng phản hồi dữ liệu lỗi C4 sang M2 chậm", "Hold")]:
        rows.append([m, q, t, g])
    grid(doc, rows)
    para(doc, "Nhóm Hold chiếm sáu trên mười ba phát hiện và tập trung ở hai loại khác "
              "nhau: chờ actor bên ngoài (IR-05, IR-10) và chờ nguồn lực nội bộ (IR-03, "
              "IR-04, IR-09). Hai loại này cần hai hướng cải tiến khác nhau, vì cửa hàng "
              "chỉ kiểm soát được loại thứ hai.")
    h2(doc, "4.7. Đề xuất cải tiến và mức độ ưu tiên")
    todo(doc, "Danh", "viết sau khi có số liệu định lượng của mục 4.2 (hạn 02/09)")

    h1(doc, "KẾT LUẬN")
    todo(doc, "Danh", "viết sau khi hoàn tất Chương 4 (hạn 03/09)")
    h1(doc, "TÀI LIỆU THAM KHẢO")
    todo(doc, "Thanh Phúc", "danh mục IEEE, tách tiếng Việt và tiếng Anh (hạn 02/09)")
    h1(doc, "PHỤ LỤC")
    h2(doc, "Phụ lục A. Hồ sơ chi tiết các quy trình")
    todo(doc, "Danh", "chèn hồ sơ C1, C3, C4 từ thư mục docs/ho-so-quy-trinh")
    h2(doc, "Phụ lục B. Bộ câu hỏi phỏng vấn")
    todo(doc, "Thanh Phúc", "24 câu chia 4 nhóm (hạn 26/08)")
    h2(doc, "Phụ lục C. Bằng chứng khảo sát")
    todo(doc, "Hưng", "bảng bấm giờ, ảnh biểu mẫu, sơ đồ mặt bằng (hạn 26/08)")


if __name__ == "__main__":
    d = build()
    chuong_3_4(d)
    d.save(OUT)
    print("Da ghi", OUT.name, "|", len(d.paragraphs), "doan,", len(d.tables), "bang")
